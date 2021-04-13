#!/usr/bin/env python3
"""
合同文档处理器 - 支持长上下文的合同审核
"""

import os
import json
from typing import List, Dict, Any
from dataclasses import dataclass
import tiktoken
import PyPDF2
from docx import Document

@dataclass
class ContractDocument:
    """合同文档数据结构"""
    content: str
    metadata: Dict[str, Any]
    token_count: int
    sections: List[Dict[str, str]]

class LongContextContractProcessor:
    """长上下文合同处理器"""
    
    def __init__(self, model_name: str = "gpt-4-turbo"):
        self.model_name = model_name
        self.encoding = tiktoken.encoding_for_model("gpt-4")
        self.max_context_length = {
            "gpt-4-turbo": 128000,
            "claude-3-sonnet": 200000,
            "gemini-1.5-pro": 1000000
        }
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """从PDF提取文本"""
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """从DOCX提取文本"""
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def count_tokens(self, text: str) -> int:
        """计算token数量"""
        return len(self.encoding.encode(text))
    
    def split_into_sections(self, text: str) -> List[Dict[str, str]]:
        """将合同分割为逻辑章节"""
        # 基于常见合同结构分割
        section_markers = [
            "第一条", "第二条", "第三条", "第四条", "第五条",
            "甲方", "乙方", "丙方",
            "1.", "2.", "3.", "4.", "5.",
            "一、", "二、", "三、", "四、", "五、",
            "条款", "协议", "附件", "补充说明"
        ]
        
        sections = []
        current_section = {"title": "开头", "content": ""}
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if any(marker in line for marker in section_markers):
                if current_section["content"]:
                    sections.append(current_section)
                current_section = {"title": line[:50], "content": line + "\n"}
            else:
                current_section["content"] += line + "\n"
        
        if current_section["content"]:
            sections.append(current_section)
        
        return sections
    
    def process_contract(self, file_path: str) -> ContractDocument:
        """处理单个合同文档"""
        # 根据文件类型提取文本
        if file_path.endswith('.pdf'):
            content = self.extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            content = self.extract_text_from_docx(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        # 计算token数量
        token_count = self.count_tokens(content)
        
        # 分割章节
        sections = self.split_into_sections(content)
        
        # 生成元数据
        metadata = {
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "token_count": token_count,
            "section_count": len(sections),
            "can_fit_in_context": token_count < self.max_context_length.get(self.model_name, 128000)
        }
        
        return ContractDocument(
            content=content,
            metadata=metadata,
            token_count=token_count,
            sections=sections
        )
    
    def combine_multiple_contracts(self, file_paths: List[str]) -> ContractDocument:
        """合并多个合同文档"""
        combined_content = ""
        combined_sections = []
        total_tokens = 0
        metadata_list = []
        
        for i, file_path in enumerate(file_paths):
            contract = self.process_contract(file_path)
            
            # 添加文档分隔符
            combined_content += f"\n\n=== 合同文档 {i+1}: {os.path.basename(file_path)} ===\n\n"
            combined_content += contract.content
            
            # 为每个章节添加文档标识
            for section in contract.sections:
                section["source_document"] = f"文档{i+1}: {os.path.basename(file_path)}"
                combined_sections.append(section)
            
            total_tokens += contract.token_count
            metadata_list.append(contract.metadata)
        
        combined_metadata = {
            "total_documents": len(file_paths),
            "total_tokens": total_tokens,
            "individual_metadata": metadata_list,
            "can_fit_in_context": total_tokens < self.max_context_length.get(self.model_name, 128000)
        }
        
        return ContractDocument(
            content=combined_content,
            metadata=combined_metadata,
            token_count=total_tokens,
            sections=combined_sections
        )

# 使用示例
if __name__ == "__main__":
    processor = LongContextContractProcessor("gemini-1.5-pro")
    
    # 处理单个合同
    contract = processor.process_contract("sample_contract.pdf")
    print(f"合同token数量: {contract.token_count}")
    print(f"是否适合单次处理: {contract.metadata['can_fit_in_context']}")
    
    # 处理多个合同
    contract_files = ["contract1.pdf", "contract2.docx", "contract3.txt"]
    combined_contract = processor.combine_multiple_contracts(contract_files)
    print(f"合并后总token数量: {combined_contract.token_count}")
